"""
知识库管理 API

提供知识库的 CRUD 操作：
- 列出知识库结构
- 上传文档到知识库
- 删除知识库文档
- 刷新知识库索引
- 搜索知识库

鉴权约定：所有写操作（upload/import/delete/refresh/ontology/*/sync/trigger）均使用
Depends(verify_internal_token)，仅内网或带 X-Internal-Token/Authorization 的请求可调用。
"""

import asyncio
import os
import re
import shutil
import threading
import time
import uuid
import json
import logging
from pathlib import Path
try:
    import fcntl
except ImportError:
    fcntl = None  # Windows 等平台无 fcntl，锁为 no-op
from datetime import datetime, timezone
from typing import List, Optional, Dict, Any
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from pydantic import BaseModel

from backend.api.deps import verify_internal_token
from backend.config.store_namespaces import NS_BOARD_PERSONAL

logger = logging.getLogger(__name__)

# 知识库根目录（与 paths 一致）
try:
    from backend.tools.base.paths import KB_PATH, ONTOLOGY_PATH, get_project_root, get_workspace_root
    KB_ROOT = KB_PATH
    ONTOLOGY_DIR = ONTOLOGY_PATH
except ImportError:
    KB_ROOT = Path(__file__).resolve().parent.parent.parent / "knowledge_base"
    ONTOLOGY_DIR = KB_ROOT / "learned" / "ontology"
    get_project_root = lambda: Path(__file__).resolve().parent.parent.parent
    get_workspace_root = lambda: get_project_root() / "tmp"

router = APIRouter(prefix="/knowledge", tags=["knowledge"])

# 后台任务引用，避免 create_task 被 GC 提前回收
_background_tasks: set = set()

# 知识库根目录的 Path 形式（用于路径解析与校验）
_KB_ROOT_PATH = Path(KB_ROOT) if not isinstance(KB_ROOT, Path) else KB_ROOT

# 允许的 team_id / user_id 字符（防路径穿越）
_ID_PATTERN = re.compile(r"^[a-zA-Z0-9_-]+$")


def _safe_error_detail(e: Exception) -> str:
    """生产环境不暴露内部异常详情，仅开发环境返回 str(e)。"""
    if os.getenv("APP_ENV", "production") == "development":
        return str(e)
    return "内部服务器错误"


def _validate_optional_id(value: Optional[str], name: str) -> None:
    """若 value 非空则校验格式，不合法则抛出 HTTPException 400。"""
    if value and not _ID_PATTERN.match(value):
        raise HTTPException(status_code=400, detail=f"{name} 不合法")


def _resolve_source_path_allowed(source_path: str) -> Path:
    """校验 source_path 在项目根或工作区内，防路径穿越，返回解析后的绝对路径。"""
    if not source_path or "\0" in source_path:
        raise HTTPException(status_code=400, detail="源路径不合法")
    raw = Path(source_path.strip().replace("\\", "/")).resolve()
    if not raw.exists() or not raw.is_dir():
        raise HTTPException(status_code=400, detail="源路径不存在或不是目录")
    try:
        proot = get_project_root().resolve()
        wroot = get_workspace_root().resolve()
    except Exception:
        proot = _KB_ROOT_PATH.resolve().parent
        wroot = proot / "tmp"
    try:
        raw.relative_to(proot)
        return raw
    except ValueError:
        pass
    try:
        raw.relative_to(wroot)
        return raw
    except ValueError:
        pass
    raise HTTPException(status_code=400, detail="源路径必须在项目根或工作区目录内")


def _resolve_path_within_kb(relative_path: str) -> Path:
    """将相对路径解析为绝对路径，并校验不逃逸知识库根目录。"""
    if not relative_path or relative_path.strip() != relative_path or relative_path in (".", "/"):
        raise HTTPException(status_code=400, detail="路径不合法")
    if "\0" in relative_path:
        raise HTTPException(status_code=400, detail="路径包含非法字符")
    relative_path = relative_path.replace("\\", "/")
    parts = Path(relative_path).parts
    if ".." in parts or (parts and parts[0] == "/"):
        raise HTTPException(status_code=400, detail="路径不合法")
    raw_path = _KB_ROOT_PATH / relative_path
    current = _KB_ROOT_PATH
    for part in parts:
        current = current / part
        if current.exists() and current.is_symlink():
            raise HTTPException(status_code=400, detail="不允许符号链接")
    resolved = raw_path.resolve()
    root_resolved = _KB_ROOT_PATH.resolve()
    try:
        resolved.relative_to(root_resolved)
    except ValueError:
        raise HTTPException(status_code=400, detail="路径不合法")
    return resolved


def _extract_workspace_thread_meta_from_source(source_path: str) -> Dict[str, Optional[str]]:
    """
    从知识库 source 路径提取 workspace/thread 维度信息。
    约定路径片段:
    - .../workspace_<workspace_id>/...
    - .../thread_<thread_id>/...
    """
    text = str(source_path or "").replace("\\", "/")
    workspace_id = None
    thread_id = None
    m_ws = re.search(r"/workspace_([a-z0-9_-]{1,64})(?:/|$)", text, flags=re.IGNORECASE)
    if m_ws:
        workspace_id = m_ws.group(1)
    m_thread = re.search(r"/thread_([a-z0-9_-]{1,64})(?:/|$)", text, flags=re.IGNORECASE)
    if m_thread:
        thread_id = m_thread.group(1)
    return {
        "workspace_id": workspace_id,
        "thread_id": thread_id,
        "source_scope": "workspace_upload" if workspace_id else "knowledge_base",
    }


# ============================================================
# 数据模型
# ============================================================

class KBDocument(BaseModel):
    """知识库文档"""
    path: str
    name: str
    type: str  # 'file' | 'directory'
    size: Optional[int] = None
    category: Optional[str] = None
    children: Optional[List["KBDocument"]] = None


class KBSearchResult(BaseModel):
    """搜索结果"""
    content: str
    source: str
    score: float
    metadata: Dict[str, Any]


class KBUploadResponse(BaseModel):
    """上传响应"""
    success: bool
    path: str
    message: str
    ontology_build_triggered: Optional[bool] = None  # True=已触发本体构建；False/None=未触发，需执行构建任务才能生成图谱


class KBRefreshResponse(BaseModel):
    """刷新响应"""
    success: bool
    documents_count: int
    chunks_count: int
    message: str


# ============================================================
# API 端点
# ============================================================

def _scan_directory(path: Path, max_depth: int = 3, current_depth: int = 0) -> List[Dict]:
    """递归扫描目录。max_depth=1 时仅返回根的直接子项（子目录的 children 为空），用于懒加载。"""
    if not path.exists() or current_depth >= max_depth:
        return []
    items = []
    try:
        entries = sorted(path.iterdir())
    except (PermissionError, FileNotFoundError, OSError) as e:
        logger.debug("_scan_directory iterdir 失败 %s: %s", path, e)
        return []
    for item in entries:
        try:
            if item.name.startswith('.') or item.is_symlink():
                continue
            if item.is_dir():
                children = _scan_directory(item, max_depth, current_depth + 1)
                items.append({
                    "path": str(item.relative_to(KB_ROOT)),
                    "name": item.name,
                    "type": "directory",
                    "children": children,
                    "count": sum(1 for c in children if c["type"] == "file") + sum(c.get("count", 0) for c in children if c["type"] == "directory"),
                })
            else:
                if item.suffix.lower() in ['.md', '.txt', '.pdf', '.docx', '.doc']:
                    try:
                        size = item.stat().st_size
                    except (FileNotFoundError, OSError, PermissionError):
                        size = 0
                    items.append({
                        "path": str(item.relative_to(KB_ROOT)),
                        "name": item.name,
                        "type": "file",
                        "size": size,
                        "extension": item.suffix.lower(),
                    })
        except (FileNotFoundError, OSError, PermissionError) as e:
            logger.debug("_scan_directory 跳过项 %s: %s", getattr(item, "name", item), e)
            continue
    return items


@router.get("/structure")
async def get_knowledge_structure(
    scope: str = "all",  # all | global | teams | users
    team_id: Optional[str] = None,
    user_id: Optional[str] = None,
    max_depth: int = 1,  # 1=仅顶层+直接子项（懒加载），3=全量递归（兼容旧版由前端传 3）
    workspace_path: Optional[str] = None,  # 前端真源 maibot_workspace_path，与当前工作区一致
) -> Dict[str, Any]:
    """
    获取知识库结构。
    max_depth=1 时只返回根及其直接子项（不递归），子级由 GET /knowledge/list 按需加载。
    workspace_path 由前端传入（与 maibot_workspace_path 一致），用于工作区真源统一。
    """
    _SCOPE_WHITELIST = ("all", "global", "teams", "users")
    if scope not in _SCOPE_WHITELIST:
        raise HTTPException(
            status_code=400,
            detail=f"非法 scope: {scope!r}，允许值: {', '.join(_SCOPE_WHITELIST)}",
        )
    try:
        depth = max(1, min(5, max_depth))
        result = {"scope": scope, "structure": []}

        if scope in ["all", "global"]:
            global_path = KB_ROOT / "global"
            if global_path.exists():
                result["structure"].append({
                    "path": "global",
                    "name": "全局知识库",
                    "type": "directory",
                    "children": _scan_directory(global_path, max_depth=depth),
                })

        if scope in ["all", "teams"]:
            teams_path = KB_ROOT / "teams"
            if teams_path.exists():
                if team_id:
                    _validate_optional_id(team_id, "team_id")
                    team_path = teams_path / team_id
                    if team_path.exists():
                        result["structure"].append({
                            "path": f"teams/{team_id}",
                            "name": f"团队: {team_id}",
                            "type": "directory",
                            "children": _scan_directory(team_path, max_depth=depth),
                        })
                else:
                    result["structure"].append({
                        "path": "teams",
                        "name": "团队知识库",
                        "type": "directory",
                        "children": _scan_directory(teams_path, max_depth=depth),
                    })

        if scope in ["all", "users"]:
            users_path = KB_ROOT / "users"
            if users_path.exists():
                if user_id:
                    _validate_optional_id(user_id, "user_id")
                    user_path = users_path / user_id
                    if user_path.exists():
                        result["structure"].append({
                            "path": f"users/{user_id}",
                            "name": f"个人: {user_id}",
                            "type": "directory",
                            "children": _scan_directory(user_path, max_depth=depth),
                        })
                else:
                    result["structure"].append({
                        "path": "users",
                        "name": "个人知识库",
                        "type": "directory",
                        "children": _scan_directory(users_path, max_depth=depth),
                    })

        return result
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("获取知识库结构失败: %s", e)
        raise HTTPException(status_code=500, detail="获取知识库结构失败")


@router.get("/tree/roots")
async def get_knowledge_tree_roots(
    scope: str = "all",
    team_id: Optional[str] = None,
    user_id: Optional[str] = None,
    workspace_path: Optional[str] = None,
) -> List[Dict[str, Any]]:
    """
    仅返回知识库根节点列表，无 children。子级由 GET /knowledge/list?path=... 按需加载。
    workspace_path 与前端工作区真源一致。
    """
    roots: List[Dict[str, Any]] = []
    if scope in ["all", "global"]:
        if (KB_ROOT / "global").exists():
            roots.append({"path": "global", "name": "全局知识库", "type": "directory"})
    if scope in ["all", "teams"]:
        if team_id:
            _validate_optional_id(team_id, "team_id")
            if (KB_ROOT / "teams" / team_id).exists():
                roots.append({"path": f"teams/{team_id}", "name": f"团队: {team_id}", "type": "directory"})
        elif (KB_ROOT / "teams").exists():
            roots.append({"path": "teams", "name": "团队知识库", "type": "directory"})
    if scope in ["all", "users"]:
        if user_id:
            _validate_optional_id(user_id, "user_id")
            if (KB_ROOT / "users" / user_id).exists():
                roots.append({"path": f"users/{user_id}", "name": f"个人: {user_id}", "type": "directory"})
        elif (KB_ROOT / "users").exists():
            roots.append({"path": "users", "name": "个人知识库", "type": "directory"})
    return roots


def _list_directory_one_level(dir_path: Path, base_relative: str) -> List[Dict[str, Any]]:
    """列出一层目录内容，返回 path/name/type/size，不递归。"""
    if not dir_path.exists() or not dir_path.is_dir():
        return []
    items = []
    try:
        for item in sorted(dir_path.iterdir()):
            if item.name.startswith('.') or item.is_symlink():
                continue
            rel = f"{base_relative}/{item.name}" if base_relative else item.name
            if item.is_dir():
                items.append({
                    "path": rel,
                    "name": item.name,
                    "type": "directory",
                })
            else:
                if item.suffix.lower() in ['.md', '.txt', '.pdf', '.docx', '.doc']:
                    items.append({
                        "path": rel,
                        "name": item.name,
                        "type": "file",
                        "size": item.stat().st_size,
                        "extension": item.suffix.lower(),
                    })
    except PermissionError as e:
        logger.debug("_scan_directory PermissionError %s: %s", path, e)
    return items


@router.get("/list")
async def list_directory(
    path: str = "",
    max_depth: int = 1,
    workspace_path: Optional[str] = None,
) -> List[Dict[str, Any]]:
    """
    列出知识库指定目录的内容（懒加载用）。
    workspace_path 与前端工作区真源一致（可选）。
    """
    path = (path or "").strip().replace("\\", "/")
    if not path:
        target = _KB_ROOT_PATH
        base_relative = ""
    else:
        target = _resolve_path_within_kb(path)
        if not target.is_dir():
            raise HTTPException(status_code=400, detail="路径不是目录")
        base_relative = path
    return _list_directory_one_level(target, base_relative)


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    target_path: str = Form(...),  # e.g., "global/domain/sales" or "users/demo-user"
    build_ontology: Optional[bool] = Form(None),  # True=上传后触发本体构建；None=按环境变量 AUTO_BUILD_ONTOLOGY_AFTER_UPLOAD
    _: None = Depends(verify_internal_token),
) -> KBUploadResponse:
    """
    上传文档到知识库。上传后仅刷新向量索引；如需生成知识图谱/本体，请传 build_ontology=True 或执行构建任务（含 ontology 步骤）/ POST /knowledge/ontology/build。
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="文件名不能为空")
    # 验证文件类型
    allowed_extensions = ['.md', '.txt', '.pdf', '.docx', '.doc']
    file_ext = Path(file.filename).suffix.lower()
    if file_ext not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件类型: {file_ext}。支持的类型: {', '.join(allowed_extensions)}"
        )
    
    # 构建目标路径（防路径穿越）
    target_dir = _resolve_path_within_kb(target_path)
    target_dir.mkdir(parents=True, exist_ok=True)
    safe_filename = Path(file.filename).name  # 只取文件名，去掉目录部分
    target_file = (target_dir / safe_filename).resolve()
    try:
        target_file.relative_to(_KB_ROOT_PATH.resolve())
    except ValueError:
        raise HTTPException(status_code=400, detail="路径不合法")
    
    # 流式写入：分块读取避免整包入内存（与 Cowork/最佳实践一致），防止大文件 OOM
    MAX_UPLOAD_SIZE = 50 * 1024 * 1024  # 50MB
    CHUNK_SIZE = 1024 * 1024  # 1MB per chunk
    try:
        total = 0
        with open(target_file, "wb") as f:
            while True:
                chunk = await file.read(CHUNK_SIZE)
                if not chunk:
                    break
                total += len(chunk)
                if total > MAX_UPLOAD_SIZE:
                    raise HTTPException(
                        status_code=413,
                        detail=f"文件过大，最大 {MAX_UPLOAD_SIZE // 1024 // 1024}MB",
                    )
                f.write(chunk)
        rebuilt = await _rebuild_knowledge_index_incremental()
        _refresh_kb_caches()
        rebuild_hint = "（索引已增量更新）" if rebuilt else "（索引重建失败，请稍后手动刷新）"
        _knowledge_metrics["upload_count"] = _knowledge_metrics.get("upload_count", 0) + 1
        ontology_triggered = False
        if _should_build_ontology_after_upload(build_ontology):
            asyncio.create_task(_trigger_ontology_build_after_upload(target_path))
            ontology_triggered = True
        msg = f"文件已上传到: {target_path}/{file.filename} {rebuild_hint}"
        if ontology_triggered:
            msg += " 已触发本体构建（后台执行），知识图谱将更新。"
        else:
            msg += " 如需生成知识图谱，请传 build_ontology=true 或执行构建任务（含 ontology 步骤）或调用 POST /knowledge/ontology/build。"
        return KBUploadResponse(
            success=True,
            path=str(target_file.relative_to(_KB_ROOT_PATH)),
            message=msg,
            ontology_build_triggered=ontology_triggered if ontology_triggered else None,
        )
    except HTTPException:
        target_file.unlink(missing_ok=True)
        raise
    except Exception as e:
        target_file.unlink(missing_ok=True)
        logger.exception("上传失败")
        raise HTTPException(status_code=500, detail="上传失败，请稍后重试")


class KBImportUrlBody(BaseModel):
    url: str
    target_path: Optional[str] = "global/imported"
    filename: Optional[str] = None


class KBImportFolderBody(BaseModel):
    source_path: str
    target_scope: str = "global/domain/sales"
    recursive: bool = True
    file_types: Optional[List[str]] = None
    build_ontology: Optional[bool] = None  # True=导入后触发本体构建；None=按环境变量 AUTO_BUILD_ONTOLOGY_AFTER_UPLOAD


class KBBuildTaskBody(BaseModel):
    source_path: Optional[str] = ""
    target_scope: str = "global/domain/sales"
    operations: Optional[List[str]] = None


@router.post("/import-url")
async def import_from_url(body: KBImportUrlBody, _: None = Depends(verify_internal_token)) -> Dict[str, Any]:
    """
    从 URL 导入资源到知识库：抓取 URL 内容并保存为文件。
    target_path 为相对知识库根目录；filename 可选，默认从 URL 或生成。
    """
    import httpx
    if not body.url or not body.url.strip():
        raise HTTPException(status_code=400, detail="url 必填")
    url = body.url.strip()
    target_path = (body.target_path or "global/imported").strip()
    target_dir = _resolve_path_within_kb(target_path)
    target_dir.mkdir(parents=True, exist_ok=True)
    filename = body.filename
    if not filename:
        try:
            from urllib.parse import urlparse, unquote
            parsed = urlparse(url)
            name = unquote(Path(parsed.path).name or "imported")
            if not name or name == ".":
                name = f"imported-{datetime.now(timezone.utc).strftime('%Y%m%d%H%M%S')}.txt"
            filename = name
        except Exception:
            filename = f"imported-{datetime.now(timezone.utc).strftime('%Y%m%d%H%M%S')}.txt"
    safe_name = Path(filename).name
    target_file = (target_dir / safe_name).resolve()
    try:
        target_file.relative_to(_KB_ROOT_PATH.resolve())
    except ValueError:
        raise HTTPException(status_code=400, detail="路径不合法")
    from backend.utils.security import is_safe_callback_url
    if not is_safe_callback_url(url):
        raise HTTPException(status_code=400, detail="URL 仅允许 https 或 localhost，且不允许内网/回环地址")
    try:
        async with httpx.AsyncClient(timeout=30.0, follow_redirects=False) as client:
            resp = await client.get(url)
            resp.raise_for_status()
            content = resp.text
    except httpx.HTTPError as e:
        logger.warning("抓取 URL 失败: %s", e)
        raise HTTPException(status_code=400, detail=f"抓取 URL 失败: {_safe_error_detail(e)}")
    try:
        target_file.write_text(content, encoding="utf-8")
        rebuilt = await _rebuild_knowledge_index_incremental()
        _refresh_kb_caches()
    except Exception as e:
        logger.exception("写入失败")
        raise HTTPException(status_code=500, detail="写入知识库失败")
    rel = str(target_file.relative_to(_KB_ROOT_PATH))
    return {
        "success": True,
        "path": rel,
        "message": f"已从 URL 导入到 {rel}",
        "index_rebuilt": rebuilt,
    }


DEFAULT_IMPORT_FILE_TYPES = [".md", ".txt", ".pdf", ".docx"]


def _validate_target_scope(target_scope: str) -> None:
    """校验 target_scope 格式，非法则抛出 ValueError。"""
    if ".." in target_scope:
        raise ValueError("target_scope 不能包含 ..")
    for part in target_scope.split("/"):
        if part and not _ID_PATTERN.match(part):
            raise ValueError("target_scope 各段仅允许字母数字、下划线、中划线")


def _copy_folder_to_kb(
    source_dir: Path,
    target_dir: Path,
    extensions: List[str],
    recursive: bool = True,
) -> tuple:
    """
    将 source_dir 下匹配扩展名的文件复制到 target_dir，保留子目录结构。
    返回 (imported_files, errors, skipped_count)。
    target_dir 须已在 KB 内（由调用方通过 _resolve_path_within_kb 得到）。
    """
    imported_files: List[str] = []
    errors: List[str] = []
    skipped_count = 0
    iterator = source_dir.rglob("*") if recursive else source_dir.iterdir()
    for fp in iterator:
        if not fp.is_file() or fp.suffix.lower() not in extensions:
            if fp.is_file():
                skipped_count += 1
            continue
        try:
            rel = fp.relative_to(source_dir)
            dest = (target_dir / rel).resolve()
            try:
                dest.relative_to(_KB_ROOT_PATH.resolve())
            except ValueError:
                errors.append(f"逃逸目标: {fp.name}")
                continue
            dest.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(fp, dest)
            imported_files.append(str(dest.relative_to(_KB_ROOT_PATH)))
        except Exception as e:
            errors.append(f"{fp.name}: {e}")
    return imported_files, errors, skipped_count


def _refresh_kb_caches() -> None:
    """清除 KnowledgeBaseManager 内存缓存，使后续检索使用最新索引。"""
    try:
        from backend.knowledge_base.manager import KnowledgeBaseManager
        KnowledgeBaseManager._global_store = None
        KnowledgeBaseManager._global_docs = []
        KnowledgeBaseManager._team_stores.clear()
        KnowledgeBaseManager._team_docs.clear()
        KnowledgeBaseManager._user_stores.clear()
        KnowledgeBaseManager._user_docs.clear()
    except Exception as e:
        logger.warning("刷新知识库缓存失败: %s", e)


_last_rebuild_at: float = 0.0
_REBUILD_DEBOUNCE_SEC = 0.1
_rebuild_lock = asyncio.Lock()

# P2 可观测：知识库/本体构建打点（上传、构建触发、成功/失败）
_knowledge_metrics: Dict[str, int] = {
    "upload_count": 0,
    "import_count": 0,
    "ontology_build_triggered": 0,
    "ontology_build_success": 0,
    "ontology_build_failure": 0,
}


def _run_ontology_build_for_scope_sync(scope: str) -> bool:
    """对指定知识库相对路径（scope）执行本体构建（同步，供 to_thread 调用）。支持 .md/.txt/.pdf/.docx/.doc。"""
    builder = _get_ontology_builder()
    if not builder or not scope:
        return False
    try:
        _validate_target_scope(scope)
    except ValueError:
        return False
    try:
        dir_path = _resolve_path_within_kb(scope)
        builder.build_from_directory(str(dir_path), recursive=True, file_types=None)
        return True
    except Exception as e:
        logger.warning("上传/导入后本体构建失败 scope=%s: %s", scope, e)
        return False


async def _trigger_ontology_build_after_upload(scope: str) -> bool:
    """上传/导入后可选触发本体构建（异步，不阻塞响应）。"""
    _knowledge_metrics["ontology_build_triggered"] = _knowledge_metrics.get("ontology_build_triggered", 0) + 1
    try:
        ok = await asyncio.to_thread(_run_ontology_build_for_scope_sync, scope)
        if ok:
            _knowledge_metrics["ontology_build_success"] = _knowledge_metrics.get("ontology_build_success", 0) + 1
        else:
            _knowledge_metrics["ontology_build_failure"] = _knowledge_metrics.get("ontology_build_failure", 0) + 1
        return ok
    except Exception as e:
        logger.warning("触发本体构建异常: %s", e)
        _knowledge_metrics["ontology_build_failure"] = _knowledge_metrics.get("ontology_build_failure", 0) + 1
        return False


def _should_build_ontology_after_upload(request_param: Optional[bool]) -> bool:
    """是否在上传/导入后自动构建本体。request_param 为请求体/查询参数，None 时用环境变量。"""
    if request_param is not None:
        return bool(request_param)
    return os.environ.get("AUTO_BUILD_ONTOLOGY_AFTER_UPLOAD", "false").strip().lower() in ("true", "1", "yes")


async def _rebuild_knowledge_index_incremental() -> bool:
    """统一增量索引重建入口；100ms 防抖，避免高频上传连续触发。"""
    import time
    async with _rebuild_lock:
        now = time.monotonic()
        if now - _last_rebuild_at < _REBUILD_DEBOUNCE_SEC:
            return True
        try:
            from backend.tools.base.embedding_tools import rebuild_index
            out = await asyncio.to_thread(rebuild_index, None, None, True, False)
            return out
        finally:
            _last_rebuild_at = time.monotonic()


@router.post("/import-folder")
async def import_folder(body: KBImportFolderBody, _: None = Depends(verify_internal_token)) -> Dict[str, Any]:
    """
    从指定文件夹批量导入到知识库：扫描 source_path 下匹配文件，复制到 knowledge_base/{target_scope}/，保留子目录结构，并触发索引刷新。
    source_path 须为项目根或工作区内的绝对路径，否则返回 400。
    """
    source_dir = _resolve_source_path_allowed(body.source_path)
    target_path = (body.target_scope or "global/domain/sales").strip().replace("\\", "/").lstrip("/")
    try:
        _validate_target_scope(target_path)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    target_dir = _resolve_path_within_kb(target_path)
    target_dir.mkdir(parents=True, exist_ok=True)
    extensions = [e.lower() if e.startswith(".") else f".{e.lower()}" for e in (body.file_types or DEFAULT_IMPORT_FILE_TYPES)]
    imported_files, errors, skipped_count = _copy_folder_to_kb(
        source_dir, target_dir, extensions, recursive=body.recursive
    )
    if imported_files:
        _knowledge_metrics["import_count"] = _knowledge_metrics.get("import_count", 0) + 1
    rebuilt = False
    if imported_files:
        rebuilt = await _rebuild_knowledge_index_incremental()
        _refresh_kb_caches()
    ontology_triggered = False
    if imported_files and _should_build_ontology_after_upload(body.build_ontology):
        ontology_triggered = await _trigger_ontology_build_after_upload(target_path)
    out = {
        "imported_count": len(imported_files),
        "skipped_count": skipped_count,
        "errors": errors[:50],
        "imported_files": imported_files[:200],
        "index_rebuilt": rebuilt,
        "ontology_build_triggered": ontology_triggered,
    }
    if imported_files:
        out["message"] = (
            "导入完成；已触发本体构建（后台执行），知识图谱将更新。"
            if ontology_triggered
            else "导入完成。如需生成知识图谱，请传 build_ontology=true 或执行构建任务（含 ontology 步骤）或调用 POST /knowledge/ontology/build。"
        )
    return out


def _board_task_update_progress(store, task_id: str, progress: int, message: Optional[str] = None) -> None:
    """更新看板任务进度（同步，供后台任务调用）。"""
    try:
        out = store.get(NS_BOARD_PERSONAL, task_id)
        if not out:
            return
        v = getattr(out, "value", out) if not isinstance(out, dict) else out
        val = dict(v) if isinstance(v, dict) else {}
        val["progress"] = max(0, min(100, progress))
        if message is not None:
            val["progress_message"] = message
        val["updated_at"] = datetime.now(timezone.utc).isoformat()
        store.put(NS_BOARD_PERSONAL, task_id, val)
    except Exception as e:
        logger.warning("更新看板任务进度失败: %s", e)


def _board_task_complete(store, task_id: str, deliverables: Optional[Dict[str, Any]] = None) -> None:
    """将看板任务标记为已完成并写入产出。"""
    try:
        out = store.get(NS_BOARD_PERSONAL, task_id)
        if not out:
            return
        v = getattr(out, "value", out) if not isinstance(out, dict) else out
        val = dict(v) if isinstance(v, dict) else {}
        val["status"] = "completed"
        val["progress"] = 100
        val["progress_message"] = "知识构建完成"
        if deliverables:
            val["deliverables"] = deliverables
        val["updated_at"] = datetime.now(timezone.utc).isoformat()
        store.put(NS_BOARD_PERSONAL, task_id, val)
        try:
            from backend.engine.tasks.task_bidding import project_board_task_status

            project_board_task_status(
                task_id=task_id,
                status="completed",
                scope="personal",
                thread_id=str(val.get("thread_id") or "") or None,
                progress=100,
                progress_message=str(val.get("progress_message") or ""),
                dispatch_state=str(val.get("dispatch_state") or "") or None,
                claimed_by=(str(val.get("claimed_by") or "") or None),
                result=str(val.get("result") or "") or None,
                source="knowledge_api_complete",
            )
        except Exception as e:
            logger.debug("project_board_task_status(complete) failed for task_id=%s: %s", task_id, e)
    except Exception as e:
        logger.warning("更新看板任务完成状态失败: %s", e)


def _board_task_fail(store, task_id: str, error: str) -> None:
    """将看板任务标记为失败。"""
    try:
        out = store.get(NS_BOARD_PERSONAL, task_id)
        if not out:
            return
        v = getattr(out, "value", out) if not isinstance(out, dict) else out
        val = dict(v) if isinstance(v, dict) else {}
        val["status"] = "failed"
        val["progress_message"] = error[:500]
        val["updated_at"] = datetime.now(timezone.utc).isoformat()
        store.put(NS_BOARD_PERSONAL, task_id, val)
        try:
            from backend.engine.tasks.task_bidding import project_board_task_status

            project_board_task_status(
                task_id=task_id,
                status="failed",
                scope="personal",
                thread_id=str(val.get("thread_id") or "") or None,
                progress=(int(val.get("progress")) if isinstance(val.get("progress"), (int, float)) else None),
                progress_message=str(val.get("progress_message") or ""),
                dispatch_state=str(val.get("dispatch_state") or "") or None,
                claimed_by=(str(val.get("claimed_by") or "") or None),
                source="knowledge_api_failed",
            )
        except Exception as e:
            logger.debug("project_board_task_status(failed) failed for task_id=%s: %s", task_id, e)
    except Exception as e:
        logger.warning("更新看板任务失败状态: %s", e)


async def _run_knowledge_build_task(task_id: str, body: KBBuildTaskBody) -> None:
    """后台执行知识构建：导入 -> 刷新索引 -> 构建本体 -> 可选 Skills。"""
    ops = body.operations or ["import", "index", "ontology"]
    source_path = (body.source_path or "").strip()
    target_scope = (body.target_scope or "global/domain/sales").strip().replace("\\", "/").lstrip("/")
    deliverables: Dict[str, Any] = {}
    try:
        from backend.engine.core.main_graph import get_sqlite_store
        store = get_sqlite_store()
        if not store:
            logger.warning("知识构建任务 Store 不可用，跳过执行")
            return
    except Exception as e:
        logger.exception("get_sqlite_store: %s", e)
        return
    deliverables["target_scope"] = target_scope
    try:
        if "import" in ops and source_path:
            await asyncio.to_thread(_board_task_update_progress, store, task_id, 10, "正在导入文件…")
            _validate_target_scope(target_scope)
            source_dir = _resolve_source_path_allowed(source_path)
            target_dir = _resolve_path_within_kb(target_scope)
            target_dir.mkdir(parents=True, exist_ok=True)
            imported, _, _ = await asyncio.to_thread(
                _copy_folder_to_kb,
                source_dir, target_dir, DEFAULT_IMPORT_FILE_TYPES, True,
            )
            deliverables["imported_files"] = imported[:100]
            deliverables["imported_count"] = len(imported)
            await asyncio.to_thread(_board_task_update_progress, store, task_id, 25, "导入完成，刷新索引…")
        else:
            if "import" in ops and not source_path:
                await asyncio.to_thread(_board_task_update_progress, store, task_id, 15, "未指定源路径，跳过导入，刷新索引…")
            await asyncio.to_thread(_board_task_update_progress, store, task_id, 25, "刷新索引…")
        if "index" in ops:
            rebuilt = await _rebuild_knowledge_index_incremental()
            _refresh_kb_caches()
            deliverables["index_rebuilt"] = bool(rebuilt)
            if rebuilt:
                await asyncio.to_thread(_board_task_update_progress, store, task_id, 50, "索引已重建，构建本体…")
            else:
                await asyncio.to_thread(_board_task_update_progress, store, task_id, 50, "索引重建失败，继续构建本体…")
        else:
            await asyncio.to_thread(_board_task_update_progress, store, task_id, 50, "未执行索引步骤，构建本体…")
        if "ontology" in ops:
            builder = _get_ontology_builder()
            if builder and target_scope:
                try:
                    dir_path = _resolve_path_within_kb(target_scope)
                    await asyncio.to_thread(
                        builder.build_from_directory,
                        str(dir_path),
                        True,
                        None,
                    )
                    deliverables["ontology_path"] = str(ONTOLOGY_DIR)
                except Exception as e:
                    logger.warning("本体构建: %s", e)
        await asyncio.to_thread(_board_task_update_progress, store, task_id, 75, "本体构建完成。")
        if "skills" in ops:
            await asyncio.to_thread(
                _board_task_update_progress,
                store,
                task_id,
                90,
                "Skills 步骤未实现：请使用现有 Skill 流程（knowledge_base/skills/ 下创建或安装 SKILL.md），或调用 knowledge-building Skill 手动从知识库沉淀步骤。",
            )
        await asyncio.to_thread(_board_task_update_progress, store, task_id, 100, "知识构建完成")
        await asyncio.to_thread(_board_task_complete, store, task_id, deliverables)
    except HTTPException as e:
        await asyncio.to_thread(_board_task_fail, store, task_id, e.detail or str(e))
    except Exception as e:
        logger.exception("知识构建任务失败: %s", e)
        await asyncio.to_thread(_board_task_fail, store, task_id, str(e))


@router.post("/build-task")
async def create_build_task(body: KBBuildTaskBody, _: None = Depends(verify_internal_token)) -> Dict[str, Any]:
    """
    创建知识构建看板任务并异步执行：导入 -> 刷新索引 -> 构建本体 -> 可选 Skills。
    operations 中 "skills" 步骤当前未实现，推荐使用现有 Skill 工作流（knowledge_base/skills/ 或 knowledge-building Skill）。
    """
    try:
        from backend.engine.core.main_graph import get_sqlite_store
        store = get_sqlite_store()
        if store is None:
            raise HTTPException(status_code=503, detail="Store 不可用")
    except Exception as e:
        logger.exception("get_sqlite_store: %s", e)
        raise HTTPException(status_code=503, detail="Store 不可用")
    task_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    val = {
        "task_id": task_id,
        "subject": "知识构建",
        "description": f"从 {body.source_path or ''} 导入并构建知识库（{body.target_scope or 'global/domain/sales'}）",
        "status": "running",
        "priority": 3,
        "scope": "personal",
        "source_channel": "local",
        "created_at": now,
        "updated_at": now,
        "progress": 0,
        "progress_message": "已创建，等待执行",
        "subtask_ids": [],
        "deliverables": None,
    }
    store.put(NS_BOARD_PERSONAL, task_id, val)
    t = asyncio.create_task(_run_knowledge_build_task(task_id, body))
    _background_tasks.add(t)
    t.add_done_callback(_background_tasks.discard)
    return {"task_id": task_id, "message": "任务已创建，可在任务面板查看进度"}


@router.delete("/document")
async def delete_document(path: str, _: None = Depends(verify_internal_token)) -> Dict[str, Any]:
    """
    删除知识库文档
    
    Args:
        path: 文档路径（相对于知识库根目录）
    
    Returns:
        删除结果
    """
    target = _resolve_path_within_kb(path)
    
    if not target.exists():
        raise HTTPException(status_code=404, detail=f"文件不存在: {path}")
    
    # 安全检查：不允许删除根目录
    if target == _KB_ROOT_PATH.resolve():
        raise HTTPException(status_code=403, detail="不能删除知识库根目录")
    
    was_dir = target.is_dir()
    try:
        if was_dir:
            shutil.rmtree(target)
        else:
            target.unlink()
    except Exception as e:
        logger.exception("删除失败")
        raise HTTPException(status_code=500, detail="删除失败，请稍后重试")
    _refresh_kb_caches()
    await _rebuild_knowledge_index_incremental()
    return {"success": True, "message": f"已删除{'目录' if was_dir else '文件'}: {path}"}


@router.post("/refresh")
async def refresh_knowledge_base(
    scope: str = "all",
    team_id: Optional[str] = None,
    user_id: Optional[str] = None,
    mode: str = "incremental",  # cache-only | incremental | full
    workspace_path: Optional[str] = None,  # 前端工作区真源，与 maibot_workspace_path 一致
    _: None = Depends(verify_internal_token),
) -> KBRefreshResponse:
    """
    刷新知识库索引（统一入口）。workspace_path 与前端工作区真源一致（可选）。
    """
    _validate_optional_id(team_id, "team_id")
    _validate_optional_id(user_id, "user_id")
    refresh_mode = str(mode or "incremental").strip().lower()
    if refresh_mode not in {"cache-only", "incremental", "full"}:
        raise HTTPException(status_code=400, detail="mode 必须是 cache-only|incremental|full")
    try:
        rebuilt = False
        if refresh_mode in {"incremental", "full"}:
            from backend.tools.base.embedding_tools import rebuild_index
            rebuilt = await asyncio.to_thread(
                rebuild_index,
                None,
                None,
                True,
                refresh_mode == "full",
            )
            if not rebuilt:
                raise HTTPException(
                    status_code=500,
                    detail=f"索引重建失败（mode={refresh_mode}）",
                )

        from backend.knowledge_base.manager import KnowledgeBaseManager
        
        # 清除缓存
        KnowledgeBaseManager._global_store = None
        KnowledgeBaseManager._global_docs = []
        
        if scope in ["all", "team", "teams"] and team_id:
            KnowledgeBaseManager._team_stores.pop(team_id, None)
            KnowledgeBaseManager._team_docs.pop(team_id, None)
        
        if scope in ["all", "user", "users"] and user_id:
            KnowledgeBaseManager._user_stores.pop(user_id, None)
            KnowledgeBaseManager._user_docs.pop(user_id, None)
        
        # 重新加载
        kb = KnowledgeBaseManager(user_id=user_id, team_id=team_id)
        
        # 统计
        docs_count = len(KnowledgeBaseManager._global_docs)
        chunks_count = 0
        if KnowledgeBaseManager._global_store:
            chunks_count = KnowledgeBaseManager._global_store.index.ntotal
        
        return KBRefreshResponse(
            success=True,
            documents_count=docs_count,
            chunks_count=chunks_count,
            message=f"知识库已刷新 (scope={scope}, mode={refresh_mode}, rebuilt={str(rebuilt).lower()})"
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=_safe_error_detail(e))


@router.get("/search")
async def search_knowledge_base(
    query: str,
    k: int = 5,
    scope: str = "all",
    team_id: Optional[str] = None,
    user_id: Optional[str] = None,
    workspace_path: Optional[str] = None,  # 前端工作区真源
) -> List[KBSearchResult]:
    """
    搜索知识库。workspace_path 与前端工作区真源一致（可选）。
    """
    _validate_optional_id(team_id, "team_id")
    _validate_optional_id(user_id, "user_id")
    try:
        # 优先使用统一的 embedding_tools（LM Studio 嵌入）
        try:
            from backend.tools.base.embedding_tools import get_knowledge_retriever_tool
            from backend.tools.base.storage_manager import get_index_manager
            from backend.engine.agent.model_manager import get_model_manager
            
            index_manager = get_index_manager()
            if index_manager.index_exists():
                from backend.tools.base.embedding_tools import get_embeddings
                embeddings = get_embeddings()
                if embeddings:
                    cache_scope = f"{scope or 'all'}_{team_id or ''}_{user_id or ''}"
                    t0 = time.perf_counter()
                    results = index_manager.search(query, embeddings, top_k=k, scope=cache_scope)
                    elapsed_ms = int((time.perf_counter() - t0) * 1000)
                    logger.info(
                        "knowledge_search_latency_ms=%s query_len=%s top_k=%s",
                        elapsed_ms, len(query), k,
                    )
                    rerank_cap = get_model_manager().get_capability_models_status().get("rerank", {})
                    rerank_status = (
                        "enabled"
                        if rerank_cap.get("enabled") and rerank_cap.get("available")
                        else ("degraded" if rerank_cap.get("enabled") else "disabled")
                    )
                    return [
                        KBSearchResult(
                            content=r.get("content", "")[:500],
                            source=r.get("source", "unknown"),
                            score=r.get("score", 0.0),
                            metadata={
                                **(r.get("metadata", {}) or {}),
                                **_extract_workspace_thread_meta_from_source(r.get("source", "")),
                                "retrieval_pipeline": {
                                    "rerank_status": rerank_status,
                                    "rerank_model": rerank_cap.get("id"),
                                    "rerank_reason": (
                                        "capability_not_available"
                                        if rerank_status == "degraded"
                                        else None
                                    ),
                                },
                            },
                        )
                        for r in results
                    ]
        except Exception as e:
            logger.warning("统一向量索引不可用，回退 KnowledgeBaseManager: %s", e)

        # 回退：使用 KnowledgeBaseManager（HuggingFace 嵌入）
        if not user_id and not team_id:
            raise HTTPException(status_code=400, detail="回退检索路径需要提供 user_id 或 team_id")
        from backend.knowledge_base.manager import KnowledgeBaseManager
        kb = KnowledgeBaseManager(
            user_id=user_id or "default",
            team_id=team_id or "default"
        )
        
        if scope == "all":
            results = kb.retrieve_multi_source(query, k=k)
        else:
            results = kb.retrieve_hybrid(query, k=k)
        
        return [
            KBSearchResult(
                content=doc.page_content[:500] + ("..." if len(doc.page_content) > 500 else ""),
                source=doc.metadata.get("source", "unknown"),
                score=doc.metadata.get("similarity_score", 0.0),
                metadata={
                    **(doc.metadata or {}),
                    **_extract_workspace_thread_meta_from_source(doc.metadata.get("source", "")),
                },
            )
            for doc in results
        ]
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=_safe_error_detail(e))


@router.get("/document")
async def get_document_content(path: str) -> Dict[str, Any]:
    """
    获取文档内容
    
    Args:
        path: 文档路径（相对于知识库根目录）
    
    Returns:
        文档内容
    """
    target = _resolve_path_within_kb(path)
    
    if not target.exists():
        raise HTTPException(status_code=404, detail=f"文件不存在: {path}")
    
    if target.is_dir():
        raise HTTPException(status_code=400, detail="不能读取目录")
    
    try:
        # 根据文件类型读取
        if target.suffix.lower() in ['.md', '.txt']:
            with open(target, 'r', encoding='utf-8') as f:
                content = f.read()
            return {
                "path": path,
                "name": target.name,
                "type": "text",
                "content": content,
            }
        elif target.suffix.lower() == '.pdf':
            # PDF 需要特殊处理
            return {
                "path": path,
                "name": target.name,
                "type": "pdf",
                "content": None,
                "message": "PDF 文件请使用专用查看器",
            }
        elif target.suffix.lower() in ['.docx', '.doc']:
            # DOCX 需要特殊处理
            try:
                import docx
                doc = docx.Document(str(target))
                content = "\n".join([p.text for p in doc.paragraphs])
                return {
                    "path": path,
                    "name": target.name,
                    "type": "docx",
                    "content": content,
                }
            except ImportError:
                return {
                    "path": path,
                    "name": target.name,
                    "type": "docx",
                    "content": None,
                    "message": "需要安装 python-docx 库",
                }
        else:
            return {
                "path": path,
                "name": target.name,
                "type": "unknown",
                "content": None,
                "message": f"不支持的文件类型: {target.suffix}",
            }
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("读取文档失败")
        raise HTTPException(status_code=500, detail="读取失败，请稍后重试")


@router.get("/document/docmap")
async def get_document_docmap(path: str) -> Dict[str, Any]:
    """
    获取文档结构（DocMap）：章节标题与行号，用于前端文档结构树与跳转。
    path: 文档路径（相对于知识库根目录）
    返回: { path, name, sections: [{ title, line, level }], keywords: [] }
    """
    target = _resolve_path_within_kb(path)
    if not target.exists() or not target.is_file():
        raise HTTPException(status_code=404, detail=f"文件不存在: {path}")
    sections = []
    keywords = []
    try:
        if target.suffix.lower() in [".md", ".txt"]:
            content = target.read_text(encoding="utf-8")
        elif target.suffix.lower() in [".docx", ".doc"]:
            try:
                import docx
                doc = docx.Document(str(target))
                content = "\n".join([p.text for p in doc.paragraphs])
            except ImportError:
                return {"path": path, "name": target.name, "sections": [], "keywords": [], "message": "需要 python-docx"}
            except Exception as e:
                return {"path": path, "name": target.name, "sections": [], "keywords": [], "message": f"docx 解析失败: {e}"}
        else:
            return {"path": path, "name": target.name, "sections": [], "keywords": []}
        lines = content.split("\n")
        for i, line in enumerate(lines, 1):
            stripped = line.strip()
            if not stripped:
                continue
            m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
            if m:
                level = len(m.group(1))
                sections.append({"title": m.group(2).strip(), "line": i, "level": level})
                continue
            m = re.match(r"^第[一二三四五六七八九十\d]+[章节条款]\s*(.+)$", stripped)
            if m:
                sections.append({"title": stripped, "line": i, "level": 1})
        if not sections and len(lines) > 0:
            sections.append({"title": target.name, "line": 1, "level": 1})
    except HTTPException:
        raise
    except Exception as e:
        logger.warning("解析 docmap 失败 %s: %s", path, e)
    return {"path": path, "name": target.name, "sections": sections, "keywords": keywords}


@router.post("/directory")
async def create_directory(path: str, _: None = Depends(verify_internal_token)) -> Dict[str, Any]:
    """
    创建知识库目录
    
    Args:
        path: 目录路径（相对于知识库根目录）
    
    Returns:
        创建结果
    """
    target = _resolve_path_within_kb(path)
    
    if target.exists():
        raise HTTPException(status_code=400, detail=f"目录已存在: {path}")
    
    try:
        target.mkdir(parents=True, exist_ok=True)
        return {"success": True, "path": path, "message": f"已创建目录: {path}"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=_safe_error_detail(e))


# ============================================================
# 知识库元数据 API（供前端展示 scopes、文档数、实体数）
# ============================================================

def _count_docs(path: Path, extensions: set) -> int:
    """统计目录下指定扩展名的文件数（不递归过深）。"""
    if not path.exists() or not path.is_dir():
        return 0
    count = 0
    try:
        for item in path.rglob("*"):
            if item.is_file() and item.suffix.lower() in extensions:
                count += 1
            if count > 10000:
                break
    except (PermissionError, OSError):
        pass
    return count


@router.get("/metadata")
async def get_knowledge_metadata() -> Dict[str, Any]:
    """
    获取知识库元数据：scopes、各 scope 文档数、本体实体/关系数。
    
    Returns:
        scopes: 各范围及文档数
        entity_count: 本体实体数
        relation_count: 本体关系数
    """
    doc_extensions = {".md", ".txt", ".pdf", ".docx", ".doc"}
    scopes = {}
    
    for scope_name, subpath in [
        ("global", "global"),
        ("teams", "teams"),
        ("users", "users"),
        ("skills", "skills"),
        ("learned", "learned"),
    ]:
        p = KB_ROOT / subpath
        scopes[scope_name] = {"path": subpath, "document_count": _count_docs(p, doc_extensions)}
    
    entity_count = 0
    relation_count = 0
    entities_path = ONTOLOGY_DIR / "entities.json"
    relations_path = ONTOLOGY_DIR / "relations.json"
    try:
        if entities_path.exists():
            data = json.loads(entities_path.read_text(encoding="utf-8"))
            entity_count = len(data.get("entities") or [])
        if relations_path.exists():
            data = json.loads(relations_path.read_text(encoding="utf-8"))
            relation_count = len(data.get("relations") or [])
    except Exception as e:
        logger.warning("读取本体文件失败: %s", e)
    
    return {
        "success": True,
        "scopes": scopes,
        "entity_count": entity_count,
        "relation_count": relation_count,
    }


@router.get("/metrics")
async def get_knowledge_metrics() -> Dict[str, Any]:
    """
    P2 可观测：知识库/本体构建打点。供审计与门禁使用。
    指标：upload_count, import_count, ontology_build_triggered, ontology_build_success, ontology_build_failure。
    构建成功率 = ontology_build_success / max(1, ontology_build_triggered)。
    """
    triggered = _knowledge_metrics.get("ontology_build_triggered", 0)
    success = _knowledge_metrics.get("ontology_build_success", 0)
    failure = _knowledge_metrics.get("ontology_build_failure", 0)
    return {
        "success": True,
        "metrics": dict(_knowledge_metrics),
        "ontology_build_success_rate": success / max(1, triggered) if triggered else None,
    }


# ============================================================
# 本体 CRUD API（知识图谱/本体人工编辑）
# ============================================================

# 串行化本体读-改-写，避免 TOCTOU 竞态（单锁覆盖 entities + relations，防死锁）
_ontology_crud_lock = asyncio.Lock()


def _ontology_file_lock(path: Path, exclusive: bool = False):
    """在支持 fcntl 的平台上对本体文件加锁（Unix）。exclusive 时返回可写句柄供 write-through。"""
    if fcntl is None:
        return None
    f = None
    try:
        path.parent.mkdir(parents=True, exist_ok=True)
        if exclusive:
            # 可覆盖写入：存在用 r+，否则用 w，便于在持锁时 write-through
            f = open(path, "r+", encoding="utf-8") if path.exists() else open(path, "w", encoding="utf-8")
        else:
            if not path.exists():
                return None  # 不存在时无需读锁
            f = open(path, "r", encoding="utf-8")
        fcntl.flock(f.fileno(), fcntl.LOCK_EX if exclusive else fcntl.LOCK_SH)
        return f
    except (AttributeError, OSError):
        if f is not None:
            try:
                f.close()
            except Exception as close_e:
                logger.debug("_ontology_file_lock close on error: %s", close_e)
        return None


def _ontology_file_unlock(f) -> None:
    """释放本体文件锁。调用方必须用 try/finally 确保所有路径（含异常）都调用本函数，避免句柄泄漏。"""
    if f is None:
        return
    try:
        if fcntl is not None:
            fcntl.flock(f.fileno(), fcntl.LOCK_UN)
    except (AttributeError, OSError, ValueError):
        pass
    try:
        f.close()
    except Exception as close_e:
        logger.debug("_ontology_file_unlock close: %s", close_e)


def _load_ontology_entities() -> tuple[list, dict]:
    """加载 entities.json，返回 (entities list, full dict)。带文件锁。"""
    path = ONTOLOGY_DIR / "entities.json"
    path.parent.mkdir(parents=True, exist_ok=True)
    if not path.exists():
        return [], {"_meta": {"description": "实体定义", "updated": None}, "entities": []}
    lock_f = _ontology_file_lock(path, exclusive=False)
    try:
        if lock_f is not None:
            lock_f.seek(0)
            data = json.loads(lock_f.read())
        else:
            data = json.loads(path.read_text(encoding="utf-8"))
        entities = data.get("entities") or []
        return entities, data
    finally:
        _ontology_file_unlock(lock_f)


def _save_ontology_entities(data: dict) -> None:
    path = ONTOLOGY_DIR / "entities.json"
    path.parent.mkdir(parents=True, exist_ok=True)
    if "_meta" not in data:
        data["_meta"] = {"description": "实体定义", "updated": datetime.now(timezone.utc).isoformat()}
    else:
        data["_meta"]["updated"] = datetime.now(timezone.utc).isoformat()
    payload = json.dumps(data, ensure_ascii=False, indent=2)
    lock_f = _ontology_file_lock(path, exclusive=True)
    try:
        if lock_f is not None:
            lock_f.seek(0)
            lock_f.truncate(0)
            lock_f.write(payload)
            lock_f.flush()
        else:
            path.write_text(payload, encoding="utf-8")
    finally:
        _ontology_file_unlock(lock_f)
    try:
        from backend.tools.base.knowledge_graph import run_ontology_backup_and_changelog
        run_ontology_backup_and_changelog(Path(ONTOLOGY_DIR) if not isinstance(ONTOLOGY_DIR, Path) else ONTOLOGY_DIR)
    except Exception as e:
        logger.warning("本体保存后备份/changelog 失败: %s", e)
    kg = _get_knowledge_graph()
    if kg:
        kg.reload()


def _load_ontology_relations() -> tuple[list, dict]:
    """加载 relations.json。带文件锁。"""
    path = ONTOLOGY_DIR / "relations.json"
    path.parent.mkdir(parents=True, exist_ok=True)
    if not path.exists():
        return [], {"_meta": {"description": "实体关系定义", "updated": None}, "relations": []}
    lock_f = _ontology_file_lock(path, exclusive=False)
    try:
        if lock_f is not None:
            lock_f.seek(0)
            data = json.loads(lock_f.read())
        else:
            data = json.loads(path.read_text(encoding="utf-8"))
        relations = data.get("relations") or []
        return relations, data
    finally:
        _ontology_file_unlock(lock_f)


def _save_ontology_relations(data: dict) -> None:
    path = ONTOLOGY_DIR / "relations.json"
    path.parent.mkdir(parents=True, exist_ok=True)
    if "_meta" not in data:
        data["_meta"] = {"description": "实体关系定义", "updated": datetime.now(timezone.utc).isoformat()}
    else:
        data["_meta"]["updated"] = datetime.now(timezone.utc).isoformat()
    payload = json.dumps(data, ensure_ascii=False, indent=2)
    lock_f = _ontology_file_lock(path, exclusive=True)
    try:
        if lock_f is not None:
            lock_f.seek(0)
            lock_f.truncate(0)
            lock_f.write(payload)
            lock_f.flush()
        else:
            path.write_text(payload, encoding="utf-8")
    finally:
        _ontology_file_unlock(lock_f)
    try:
        from backend.tools.base.knowledge_graph import run_ontology_backup_and_changelog
        run_ontology_backup_and_changelog(Path(ONTOLOGY_DIR) if not isinstance(ONTOLOGY_DIR, Path) else ONTOLOGY_DIR)
    except Exception as e:
        logger.warning("本体保存后备份/changelog 失败: %s", e)
    kg = _get_knowledge_graph()
    if kg:
        kg.reload()


@router.get("/ontology/entities")
async def list_ontology_entities() -> Dict[str, Any]:
    """列出本体实体。"""
    entities, _ = _load_ontology_entities()
    return {"success": True, "entities": entities, "total": len(entities)}


@router.get("/ontology/entities/{entity_id}")
async def get_ontology_entity(entity_id: str) -> Dict[str, Any]:
    """获取单个实体（按 id 或 name 匹配）。"""
    entities, _ = _load_ontology_entities()
    for e in entities:
        if isinstance(e, dict) and (e.get("id") == entity_id or e.get("name") == entity_id):
            return {"success": True, "entity": e}
    raise HTTPException(status_code=404, detail=f"实体不存在: {entity_id}")


@router.post("/ontology/entities")
async def create_ontology_entity(entity: Dict[str, Any], _: None = Depends(verify_internal_token)) -> Dict[str, Any]:
    """新增实体。"""
    async with _ontology_crud_lock:
        entities, data = _load_ontology_entities()
        entity_id = entity.get("id") or entity.get("name") or str(uuid.uuid4())[:8]
        entity["id"] = entity_id
        entities.append(entity)
        data["entities"] = entities
        _save_ontology_entities(data)
    return {"success": True, "entity": entity}


@router.put("/ontology/entities/{entity_id}")
async def update_ontology_entity(entity_id: str, entity: Dict[str, Any], _: None = Depends(verify_internal_token)) -> Dict[str, Any]:
    """更新实体。"""
    async with _ontology_crud_lock:
        entities, data = _load_ontology_entities()
        for i, e in enumerate(entities):
            if isinstance(e, dict) and (e.get("id") == entity_id or e.get("name") == entity_id):
                entity["id"] = entity_id
                entities[i] = entity
                data["entities"] = entities
                _save_ontology_entities(data)
                return {"success": True, "entity": entity}
    raise HTTPException(status_code=404, detail=f"实体不存在: {entity_id}")


@router.delete("/ontology/entities/{entity_id}")
async def delete_ontology_entity(entity_id: str, _: None = Depends(verify_internal_token)) -> Dict[str, Any]:
    """删除实体。"""
    async with _ontology_crud_lock:
        entities, data = _load_ontology_entities()
        new_entities = [e for e in entities if isinstance(e, dict) and e.get("id") != entity_id and e.get("name") != entity_id]
        if len(new_entities) == len(entities):
            raise HTTPException(status_code=404, detail=f"实体不存在: {entity_id}")
        data["entities"] = new_entities
        _save_ontology_entities(data)
    return {"success": True, "message": f"已删除实体: {entity_id}"}


@router.get("/ontology/relations")
async def list_ontology_relations() -> Dict[str, Any]:
    """列出本体关系。"""
    relations, _ = _load_ontology_relations()
    return {"success": True, "relations": relations, "total": len(relations)}


@router.post("/ontology/relations")
async def create_ontology_relation(relation: Dict[str, Any], _: None = Depends(verify_internal_token)) -> Dict[str, Any]:
    """新增关系。"""
    async with _ontology_crud_lock:
        relations, data = _load_ontology_relations()
        relations.append(relation)
        data["relations"] = relations
        _save_ontology_relations(data)
    return {"success": True, "relation": relation}


@router.delete("/ontology/relations/{index}")
async def delete_ontology_relation(index: int, _: None = Depends(verify_internal_token)) -> Dict[str, Any]:
    """按索引删除关系。"""
    async with _ontology_crud_lock:
        relations, data = _load_ontology_relations()
        if index < 0 or index >= len(relations):
            raise HTTPException(status_code=404, detail="关系索引无效")
        relations.pop(index)
        data["relations"] = relations
        _save_ontology_relations(data)
    return {"success": True, "message": f"已删除关系索引: {index}"}


@router.post("/ontology/import")
async def import_ontology_batch(body: Dict[str, Any], _: None = Depends(verify_internal_token)) -> Dict[str, Any]:
    """批量导入实体与关系。body: { entities: [{name?, type?, id?}], relations: [{source, target, type?}] }"""
    entities_in = body.get("entities") or []
    relations_in = body.get("relations") or []
    if not entities_in and not relations_in:
        raise HTTPException(status_code=400, detail="entities 或 relations 至少一项非空")
    added_entities = 0
    added_relations = 0
    errors = []
    async with _ontology_crud_lock:
        entities, data_ent = _load_ontology_entities()
        relations, data_rel = _load_ontology_relations()
        for i, e in enumerate(entities_in):
            if not isinstance(e, dict):
                errors.append(f"entities[{i}] 非对象")
                continue
            eid = e.get("id") or e.get("name") or f"import_{uuid.uuid4().hex[:8]}"
            e["id"] = eid
            entities.append(e)
            added_entities += 1
        if added_entities:
            data_ent["entities"] = entities
            _save_ontology_entities(data_ent)
        for i, r in enumerate(relations_in):
            if not isinstance(r, dict):
                errors.append(f"relations[{i}] 非对象")
                continue
            src, tgt = r.get("source"), r.get("target")
            if not src or not tgt:
                errors.append(f"relations[{i}] 缺少 source 或 target")
                continue
            relations.append({
                "source": src, "target": tgt, "type": r.get("type"),
                "subject_id": src, "object_id": tgt, "predicate": r.get("type", "requires"),
                "id": f"rel_{uuid.uuid4().hex[:8]}",
            })
            added_relations += 1
        if added_relations:
            data_rel["relations"] = relations
            _save_ontology_relations(data_rel)
    return {
        "success": True,
        "entities_added": added_entities,
        "relations_added": added_relations,
        "errors": errors[:20],
    }


# ============================================================
# 知识图谱可视化 API
# ============================================================

def _get_ontology_builder():
    """获取本体构建器（用于导出图谱数据）"""
    try:
        from backend.tools.base.ontology_builder import OntologyBuilder
        return OntologyBuilder()
    except ImportError as e:
        logger.warning("OntologyBuilder 不可用: %s", e)
        return None


def _get_knowledge_graph():
    """获取知识图谱实例"""
    try:
        from backend.tools.base.knowledge_graph import get_knowledge_graph
        return get_knowledge_graph()
    except ImportError:
        return None


@router.get("/graph/data")
async def get_graph_data(
    limit: int = 500,
    entity_type: Optional[str] = None,
    relation_type: Optional[str] = None,
    min_confidence: float = 0.0,
) -> Dict[str, Any]:
    """获取图谱可视化数据（节点+边）"""
    try:
        builder = _get_ontology_builder()
        if builder is None:
            return {"nodes": [], "edges": [], "stats": {}}
        data = builder.export_for_visualization(
            limit=limit,
            entity_type=entity_type,
            relation_type=relation_type,
            min_confidence=min_confidence,
        )
        return data
    except Exception as e:
        logger.exception("get_graph_data 失败: %s", e)
        raise HTTPException(status_code=500, detail="获取图谱数据失败")


@router.get("/graph/subgraph/{entity_id}")
async def get_subgraph(
    entity_id: str,
    depth: int = 2,
    max_nodes: int = 100,
) -> Dict[str, Any]:
    """获取以指定实体为中心的子图"""
    try:
        builder = _get_ontology_builder()
        if builder is None:
            return {"nodes": [], "edges": [], "stats": {}}
        return builder.export_subgraph(entity_id, depth=depth, max_nodes=max_nodes)
    except Exception as e:
        logger.exception("get_subgraph 失败: %s", e)
        raise HTTPException(status_code=500, detail="获取子图失败")


@router.get("/graph/stats")
async def get_graph_stats() -> Dict[str, Any]:
    """获取图谱统计信息"""
    try:
        kg = _get_knowledge_graph()
        if kg is None:
            return {
                "total_entities": 0,
                "total_relations": 0,
                "entities_by_type": {},
                "relations_by_type": {},
                "top_connected_entities": [],
                "isolated_entities_count": 0,
            }
        stats = kg.get_stats()
        connected = set()
        for rel in kg._relations.values():
            connected.add(rel.subject_id)
            connected.add(rel.object_id)
        isolated = sum(1 for eid in kg._entities if eid not in connected)
        by_connections = []
        for eid, entity in kg._entities.items():
            out = len(kg._relations_by_subject.get(eid, set()))
            inc = len(kg._relations_by_object.get(eid, set()))
            by_connections.append({"id": eid, "name": entity.name, "connections": out + inc})
        by_connections.sort(key=lambda x: -x["connections"])
        return {
            "total_entities": stats["total_entities"],
            "total_relations": stats["total_relations"],
            "entities_by_type": stats.get("entity_types", {}),
            "relations_by_type": stats.get("relation_types", {}),
            "top_connected_entities": by_connections[:20],
            "isolated_entities_count": isolated,
        }
    except Exception as e:
        logger.exception("get_graph_stats 失败: %s", e)
        raise HTTPException(status_code=500, detail="获取图谱统计失败")


@router.get("/graph/search")
async def search_graph(
    q: str,
    entity_type: Optional[str] = None,
    limit: int = 20,
) -> Dict[str, Any]:
    """在图谱中搜索实体及其邻域"""
    kg = _get_knowledge_graph()
    if kg is None or not q.strip():
        return {"entities": [], "total": 0}
    try:
        from backend.tools.base.knowledge_graph import EntityType
        et = EntityType(entity_type) if entity_type else None
    except (ValueError, TypeError):
        et = None
    entities = kg.find_entities(name=q, entity_type=et)
    if not entities:
        entities = kg.find_entities(name=q)
    result = []
    for e in entities[:limit]:
        result.append({
            "id": e.id,
            "name": e.name,
            "type": e.entity_type.value,
            "properties": e.properties,
        })
    return {"entities": result, "total": len(result)}


@router.get("/graph/entity/{entity_id}/neighbors")
async def get_entity_neighbors(
    entity_id: str,
    relation_types: Optional[str] = None,
    direction: str = "both",
) -> Dict[str, Any]:
    """获取实体的邻居节点"""
    try:
        builder = _get_ontology_builder()
        if builder is None:
            return {"nodes": [], "edges": [], "stats": {}}
        kg = builder.knowledge_graph
        if entity_id not in kg._entities:
            return {"nodes": [], "edges": [], "stats": {}}
        rel_types = [s.strip() for s in (relation_types or "").split(",") if s.strip()]
        try:
            from backend.tools.base.knowledge_graph import RelationType
            rtypes = [RelationType(rt) for rt in rel_types] if rel_types else None
        except (ValueError, TypeError):
            rtypes = None
        neighbors = kg.get_related_entities(entity_id, relation_types=rtypes, direction=direction, max_depth=1)
        node_ids = {entity_id}
        for ent, rel, _ in neighbors:
            node_ids.add(ent.id)
        def _entity_props(e):
            p = getattr(e, "properties", None)
            return p if isinstance(p, dict) else {}
        nodes = [{"id": eid, "label": kg._entities[eid].name, "type": kg._entities[eid].entity_type.value, "properties": _entity_props(kg._entities[eid])} for eid in node_ids if eid in kg._entities]
        edges = []
        for eid in node_ids:
            for r in kg.get_relations(subject_id=eid) + kg.get_relations(object_id=eid):
                if r.subject_id in node_ids and r.object_id in node_ids:
                    edges.append({"id": r.id, "source": r.subject_id, "target": r.object_id, "label": r.predicate.value})
        return {"nodes": nodes, "edges": edges, "stats": {"totalEntities": len(nodes), "totalRelations": len(edges)}}
    except Exception as e:
        logger.exception("get_entity_neighbors 失败: %s", e)
        raise HTTPException(status_code=500, detail="获取实体邻居失败")


# ============================================================
# 本体构建 API
# ============================================================

_build_tasks: Dict[str, Dict[str, Any]] = {}
_build_tasks_lock = threading.Lock()
_BUILD_TASKS_MAX = 100


def _prune_build_tasks() -> None:
    """保留最近 _BUILD_TASKS_MAX 条构建任务，避免无限增长。调用方需持有 _build_tasks_lock。"""
    while len(_build_tasks) >= _BUILD_TASKS_MAX:
        oldest = next(iter(_build_tasks))
        del _build_tasks[oldest]


@router.post("/ontology/build")
async def build_ontology(body: Dict[str, Any], _: None = Depends(verify_internal_token)) -> Dict[str, Any]:
    """触发本体构建。body: directory, domain, use_llm, recursive"""
    directory = body.get("directory", "")
    domain = body.get("domain", "bidding")
    use_llm = body.get("use_llm", False)
    recursive = body.get("recursive", True)
    if not directory or not str(directory).strip():
        raise HTTPException(status_code=400, detail="directory 必填")
    try:
        dir_path = _resolve_path_within_kb(str(directory).lstrip("/").strip())
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("ontology build path 解析失败")
        raise HTTPException(status_code=400, detail="路径不合法")
    task_id = str(uuid.uuid4())[:8]
    builder = _get_ontology_builder()
    if builder is None:
        raise HTTPException(status_code=503, detail="OntologyBuilder 不可用")
    try:
        result = builder.build_from_directory(
            str(dir_path),
            recursive=recursive,
            file_types=None,
        )
        with _build_tasks_lock:
            _prune_build_tasks()
            _build_tasks[task_id] = {"status": "completed", "task_id": task_id, "stats": result.get("stats", {})}
        return {"task_id": task_id, "status": "completed", "stats": result.get("stats", {})}
    except Exception as e:
        logger.exception("本体构建失败")
        with _build_tasks_lock:
            _prune_build_tasks()
            _build_tasks[task_id] = {"status": "failed", "task_id": task_id, "error": str(e)}
        raise HTTPException(status_code=500, detail="本体构建失败，请稍后重试")


@router.get("/ontology/build/status/{task_id}")
async def get_build_status(task_id: str) -> Dict[str, Any]:
    """查询构建进度"""
    with _build_tasks_lock:
        if task_id not in _build_tasks:
            return {"task_id": task_id, "status": "unknown", "stats": {}}
        return _build_tasks[task_id]


@router.post("/ontology/validate")
async def validate_ontology(_: None = Depends(verify_internal_token)) -> Dict[str, Any]:
    """验证本体一致性"""
    builder = _get_ontology_builder()
    if builder is None:
        return {"valid": True, "issues": [], "stats": {}}
    return builder.validate_ontology()


@router.get("/ontology/schema")
async def get_ontology_schema() -> Dict[str, Any]:
    """获取本体模式定义"""
    try:
        from backend.tools.base.ontology_builder import OntologySchema
        schema = OntologySchema()
        return schema._data
    except Exception as e:
        logger.warning("获取 schema 失败: %s", e)
        return {"entity_types": {}, "relation_types": {}, "domain": "general"}


@router.put("/ontology/schema")
async def update_ontology_schema(schema: Dict[str, Any], _: None = Depends(verify_internal_token)) -> Dict[str, Any]:
    """更新本体模式定义"""
    try:
        path = ONTOLOGY_DIR / "schema.json"
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(json.dumps(schema, ensure_ascii=False, indent=2), encoding="utf-8")
        return {"success": True}
    except Exception as e:
        logger.exception("update_ontology_schema 失败: %s", e)
        raise HTTPException(status_code=500, detail="更新本体模式失败")


# ============================================================
# 知识库云端同步（基于 LangGraph Store namespace + TTL）
# ============================================================

@router.get("/sync/status")
async def knowledge_sync_status(
    user_id: str = "default",
    domain: str = "bidding",
) -> Dict[str, Any]:
    """获取指定用户、领域的同步状态（last_sync_ts, expired, cached）。"""
    try:
        from backend.engine.core.main_graph import get_sqlite_store
        from backend.tools.base.knowledge_sync import get_sync_status
        store = get_sqlite_store()
        if store is None:
            return {"last_sync_ts": None, "cloud_version": None, "expired": True, "cached": False}
        return get_sync_status(store, user_id, domain)
    except Exception as e:
        logger.warning("knowledge_sync_status failed: %s", e)
        return {"last_sync_ts": None, "cloud_version": None, "expired": True, "cached": False}


@router.post("/sync/trigger")
async def knowledge_sync_trigger(body: Dict[str, Any], _: None = Depends(verify_internal_token)) -> Dict[str, Any]:
    """触发指定用户、领域的云端同步。body: { "user_id": str, "domain": str }。"""
    user_id = (body.get("user_id") or "default").strip()
    domain = (body.get("domain") or "bidding").strip()
    try:
        from backend.engine.core.main_graph import get_sqlite_store
        from backend.tools.base.knowledge_sync import sync_domain
        store = get_sqlite_store()
        if store is None:
            return {"success": False, "message": "Store 不可用", "entries_count": 0}
        return sync_domain(store, user_id, domain)
    except Exception as e:
        logger.exception("knowledge_sync_trigger failed")
        return {"success": False, "message": str(e), "entries_count": 0}


# ============================================================
# 导出
# ============================================================

__all__ = ["router"]
